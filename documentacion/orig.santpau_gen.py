import sys
import os
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from datetime import date

def insert_paragraph_after(paragraph, text=None, style=None):
    """
    Inserta un párrafo nuevo exactamente después del párrafo dado
    usando manipulación directa del árbol XML de Word.
    """
    new_p_element = OxmlElement('w:p')
    paragraph._element.addnext(new_p_element)
    new_para = Paragraph(new_p_element, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    else:
        new_para.style = paragraph.style
    return new_para

def smart_replace(paragraph, key, value):
    """
    Busca placeholders y maneja saltos de línea inyectando nuevos párrafos
    con estilo 'Normal' para no romper la estética corporativa.
    """
    placeholder_double = f"{{{{{key}}}}}"
    placeholder_tag = f"<{key}>"
    
    target = None
    if placeholder_double in paragraph.text:
        target = placeholder_double
    elif placeholder_tag in paragraph.text:
        target = placeholder_tag

    if target:
        text_value = str(value)
        if '\n' in text_value:
            lines = text_value.split('\n')
            # La primera línea reemplaza al placeholder en el párrafo actual
            paragraph.text = paragraph.text.replace(target, lines[0])
            
            # Las siguientes líneas se crean como párrafos nuevos
            current_p = paragraph
            for line in lines[1:]:
                # Solo insertamos si la línea tiene contenido para evitar huecos feos
                if line.strip() or line == "": 
                    new_p = insert_paragraph_after(current_p, line, style='Normal')
                    current_p = new_p
        else:
            paragraph.text = paragraph.text.replace(target, text_value)

def process_paragraphs(paragraphs, reemplazos):
    for p in paragraphs:
        for key, value in reemplazos.items():
            smart_replace(p, key, value)

def replace_placeholders(doc, reemplazos):
    # 1. Procesar el cuerpo del documento
    process_paragraphs(doc.paragraphs, reemplazos)
    
    # 2. Procesar todas las tablas (celda por celda)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, reemplazos)
    
    # 3. Procesar Encabezados y Pies de página (Crucial para el "Documento Íntegro")
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs, reemplazos)
        process_paragraphs(section.footer.paragraphs, reemplazos)
        # También procesar tablas dentro de headers/footers si las hay
        for header_table in section.header.tables:
            for row in header_table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs, reemplazos)

def fill_fixed_tables(doc, tablas_data):
    """
    Rellena las tablas de control de la primera página (Sant Pau Standard).
    """
    # Tabla 0: Información del documento
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        try:
            t0.cell(0, 1).text = tablas_data.get('nom_doc', '')
            t0.cell(1, 1).text = str(date.today())
            t0.cell(3, 1).text = tablas_data.get('version', '1.0')
        except IndexError: pass

    # Tabla 1: Control de cambios
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if len(t1.rows) > 1:
            r = t1.rows[1].cells
            r[0].text = tablas_data.get('version', '1.0')
            r[1].text = str(date.today())
            r[2].text = tablas_data.get('cambio_desc', '')

def main():
    if len(sys.argv) < 2:
        print("Uso: python3 santpau_gen.py NombreProyecto")
        sys.exit(1)

    try:
        with open('datos.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print("❌ Error: No se encuentra datos.json en el directorio actual.")
        sys.exit(1)

    template_name = data.get('template', 'DSI-COX-PL-Plantilla Word genèrica v0.2.docx')
    template_path = os.path.join('plantillas', template_name)
    
    if not os.path.exists(template_path):
        print(f"❌ Error: La plantilla no existe en: {template_path}")
        sys.exit(1)

    output_dir = sys.argv[1]
    if not os.path.exists(output_dir): os.makedirs(output_dir)
    
    doc = Document(template_path)
    
    # Ejecutar el motor de reemplazo
    fill_fixed_tables(doc, data['tablas'])
    replace_placeholders(doc, data['reemplazos'])

    output_path = os.path.join(output_dir, f"{sys.argv[1]}.docx")
    doc.save(output_path)
    print(f"✅ Documento íntegro generado con éxito en: {output_path}")

if __name__ == "__main__":
    main()
