#!/usr/bin/env python3
import sys
import os
import json
from datetime import date
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

def insert_paragraph_after(paragraph, text='', style='Normal'):
    """Inserta un párrafo después del actual utilizando manipulación de XML nativa."""
    new_p_elm = OxmlElement('w:p')
    paragraph._element.addnext(new_p_elm)
    new_para = Paragraph(new_p_elm, paragraph._parent)
    if text:
        new_para.text = text
    if style:
        new_para.style = style
    return new_para

def smart_replace(paragraph, key, value):
    """
    Reemplaza placeholders manejando {{CLAVE}} y <CLAVE>,
    soportando valores multilínea sin romper formatos.
    """
    placeholder_double = f"{{{{{key}}}}}"
    placeholder_tag = f"<{key}>"
    
    target = None
    if placeholder_double in paragraph.text:
        target = placeholder_double
    elif placeholder_tag in paragraph.text:
        target = placeholder_tag

    if target:
        if '\n' in str(value):
            lines = str(value).split('\n')
            # La primera línea reemplaza al placeholder
            paragraph.text = paragraph.text.replace(target, lines[0])
            current_p = paragraph
            for line in lines[1:]:
                # Insertar siguientes líneas con estilo Normal
                current_p = insert_paragraph_after(current_p, line, style='Normal')
        else:
            paragraph.text = paragraph.text.replace(target, str(value))

def process_paragraphs(paragraphs, reemplazos):
    """Aplica smart_replace en una lista de párrafos."""
    for p in paragraphs:
        for key, value in reemplazos.items():
            smart_replace(p, key, value)

def process_table(table, reemplazos):
    """Recorre celdas y párrafos de una tabla buscando placeholders."""
    for row in table.rows:
        for cell in row.cells:
            process_paragraphs(cell.paragraphs, reemplazos)

def replace_placeholders(doc, reemplazos):
    """Recorre Cuerpo, Tablas, Encabezados y Pies de página (incluyendo sus tablas)."""
    # 1. Cuerpo principal
    process_paragraphs(doc.paragraphs, reemplazos)
    
    # 2. Tablas del cuerpo
    for table in doc.tables:
        process_table(table, reemplazos)
        
    # 3. Encabezados y Pies de página (Headers & Footers)
    for section in doc.sections:
        # Párrafos de Header y Footer
        process_paragraphs(section.header.paragraphs, reemplazos)
        process_paragraphs(section.footer.paragraphs, reemplazos)
        
        # Tablas dentro del Header y Footer (Aquí se ubican TITULO_DOC y FECHA)
        for h_table in section.header.tables:
            process_table(h_table, reemplazos)
        for f_table in section.footer.tables:
            process_table(f_table, reemplazos)

def fill_fixed_tables(doc, tablas_data):
    """Rellena la información fija de las tablas iniciales de la plantilla."""
    # Tabla 0: Información del documento
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        try:
            t0.cell(0, 1).text = tablas_data.get('nom_doc', '')
            t0.cell(1, 1).text = str(date.today())
            t0.cell(3, 1).text = tablas_data.get('version', '1.0')
        except IndexError:
            pass

    # Tabla 1: Control de cambios
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if len(t1.rows) > 1:
            r = t1.rows[1].cells
            try:
                r[0].text = tablas_data.get('version', '1.0')
                r[1].text = str(date.today())
                r[2].text = tablas_data.get('cambio_desc', '')
            except IndexError:
                pass

def main():
    if len(sys.argv) < 2:
        print("Uso: python3.11 santpau_gen.py NombreProyecto")
        sys.exit(1)

    try:
        with open('datos.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print("❌ Error: No se encuentra datos.json")
        sys.exit(1)

    # Búsqueda de la plantilla (en carpeta 'plantillas' o ruta directa)
    template_setting = data.get('template', 'DSI-COX-PL-Plantilla procediment Word v0.2.docx')
    if os.path.exists(os.path.join('plantillas', template_setting)):
        template_path = os.path.join('plantillas', template_setting)
    elif os.path.exists(template_setting):
        template_path = template_setting
    else:
        print(f"❌ Error: No se encuentra la plantilla base '{template_setting}'")
        sys.exit(1)

    output_dir = sys.argv[1]
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document(template_path)

    # 1. Relleno directo de celdas en tablas estructurales
    fill_fixed_tables(doc, data['tablas'])

    # 2. Inyección de placeholders en todo el documento (cuerpo + headers/footers + tablas)
    replace_placeholders(doc, data['reemplazos'])

    output_path = os.path.join(output_dir, f"{sys.argv[1]}.docx")
    doc.save(output_path)
    print(f"✅ Documento generado con éxito en: {output_path}")

if __name__ == "__main__":
    main()
